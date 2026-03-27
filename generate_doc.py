from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

doc = Document()

# ─── Helper functions ──────────────────────────────────────────────────────────

def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return h

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()

def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_page_break()

# ─── Cover Page ────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ShadowNet Nexus v5.0")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("Unified Live Forensic Intelligence System")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.add_run("Technical Documentation — Reference for Presentation").italic = True

doc.add_paragraph()
doc.add_paragraph()

# ─── 1. Abstract ───────────────────────────────────────────────────────────────
page_break(doc)
set_heading(doc, "1. Abstract", 1, (0, 120, 212))

add_body(doc,
    "ShadowNet Nexus v5.0 is a real-time, AI-powered digital forensic intelligence system "
    "designed to detect, capture, and analyze malicious activity on a live running computer "
    "before an attacker can cover their tracks. Unlike traditional antivirus or endpoint "
    "security tools that rely on known virus signatures, ShadowNet works by observing the "
    "fundamental behavior of the operating system itself — reading raw disk structures, "
    "measuring the mathematical randomness of data on the hard drive, analyzing the rhythmic "
    "timing of human input, and intercepting every command that a running process executes."
)
add_body(doc,
    "The system pairs these hardware-level sensors with Google's Gemini 2.5 Flash AI model, "
    "which interprets the observed behavior in plain language and determines whether an attack "
    "is taking place. The result is a complete, self-healing forensic platform that functions "
    "around the clock and produces structured evidence reports automatically."
)

# ─── 2. Problem Statement ──────────────────────────────────────────────────────
page_break(doc)
set_heading(doc, "2. The Problem Being Solved", 1, (0, 120, 212))

add_body(doc,
    "Modern attackers use a set of techniques that completely defeat traditional security tools:"
)

problems = [
    ("Living off the Land (LotL)", 
     "Attackers use legitimate Windows tools already on the computer — like PowerShell, "
     "wevtutil, and vssadmin — so that no new malicious file is ever written to the disk. "
     "Standard antivirus has nothing suspicious to flag."),
    ("Timestomping", 
     "Malware alters the visible file creation date on a file to make it look like a legitimate "
     "old system file, hiding it from timeline-based investigations."),
    ("Volume Shadow Copy Deletion", 
     "Ransomware silently removes all Windows backup snapshots before encrypting files, "
     "making recovery nearly impossible."),
    ("High-Entropy Encryption / Wiping",
     "Ransomware and data wiper programs overwrite files with mathematically random data or "
     "all-zeros, making the original data irrecoverable. Standard tools cannot detect this "
     "until it is too late."),
    ("Automated Bot / Keylogger Input",
     "Automated scripts that steal passwords or perform actions on behalf of a hacker "
     "produce input patterns that are physically impossible for a human — inhuman speed, "
     "no variation. Standard systems cannot tell the difference."),
    ("Log File Destruction",
     "Highly skilled attackers delete Windows Event Logs and NTFS change journals to erase "
     "the record of what happened on a system."),
]

for title_txt, desc in problems:
    p = doc.add_paragraph()
    run = p.add_run(f"{title_txt}: ")
    run.bold = True
    p.add_run(desc)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)

# ─── 3. System Architecture ──────────────────────────────────────────────────
page_break(doc)
set_heading(doc, "3. System Architecture", 1, (0, 120, 212))

add_body(doc,
    "ShadowNet Nexus v5.0 is built as a layered, multi-threaded Python application. "
    "All modules run concurrently and report their findings through a central event bridge "
    "to a shared incident reporting and AI analysis pipeline."
)

set_heading(doc, "3.1 High-Level Layer Diagram", 2)
layers = [
    ("Layer 0 – Hardware / OS Kernel", 
     "Raw disk handle (CreateFileW), Windows Message Queue (Low-Level Keyboard Hook), WMI process events"),
    ("Layer 1 – V5 Sensor Modules",
     "NTFS Artifact Parser, Entropy Map Scanner, Real-time Behavior Monitor, File Integrity Monitor"),
    ("Layer 2 – V4 Detection Engine",
     "Process Monitor (WMI + psutil), V4 Behavioral Monitor (Simulation + AI), Alert Deduplication"),
    ("Layer 3 – AI Analysis Core",
     "Gemini Command Analyzer, Circuit Breaker, Rate Limiter, Keyword Fallback Engine"),
    ("Layer 4 – Incident Response",
     "Emergency Snapshot Engine, Incident Report Generator, Evidence Chain Signer, Proactive Evidence Collector"),
    ("Layer 5 – Resilience & Reporting",
     "Thread Watchdog, Forensic Narrative Engine (LLM Synthesis), Structured Hash-Chain Logger"),
    ("Layer 6 – Visibility",
     "Live Health Dashboard (FastAPI), Health API, SIEM Integration Hooks"),
]
add_table(doc, ["Layer", "Components"], layers)

# ─── 4. Detection Modules (Deep Dive) ─────────────────────────────────────────
page_break(doc)
set_heading(doc, "4. Detection Modules — Deep Dive", 1, (0, 120, 212))

# 4.1 Process Monitor
set_heading(doc, "4.1 Process Monitor (V4)", 2)
add_body(doc,
    "File: core/process_monitor.py"
)
add_body(doc,
    "The Process Monitor runs two parallel detection strategies. On Windows, it uses the "
    "WMI (Windows Management Instrumentation) event subscription API to receive an instant "
    "notification the moment any new process is created on the system. Simultaneously, "
    "it runs a fast psutil polling loop that checks for new PIDs every 10 milliseconds "
    "as a backup."
)
add_body(doc, "When a new process is detected:")
for step in [
    "The full command-line argument string is extracted from the process.",
    "It is checked against a configurable list of suspicious keywords (vssadmin, mimikatz, wevtutil, etc.).",
    "If suspicious, the command is immediately sent to the Gemini AI Analyzer for intent analysis.",
    "The result is forwarded to the Incident Response pipeline."
]:
    add_bullet(doc, step)

# 4.2 NTFS Parser
set_heading(doc, "4.2 NTFS Artifact Parser (V5)", 2)
add_body(doc, "File: core/ntfs_artifact_parser.py  |  774 lines")
add_body(doc,
    "This is the most technically advanced module in the system. It completely bypasses "
    "the Windows operating system's file API and opens a raw binary handle directly to the "
    "physical volume (e.g., \\\\.\\C:) using the CreateFileW kernel call. This means no "
    "rootkit or anti-forensic tool running in user-mode can hide from it."
)
add_body(doc, "What it reads and why it matters:")
ntfs_points = [
    ("$MFT (Master File Table)", 
     "Every file that exists or has ever existed on the volume has a record here. "
     "The MFT contains two independent timestamps for each file: $STANDARD_INFORMATION ($SI) "
     "which is user-accessible, and $FILE_NAME ($FN) which is a kernel-managed record "
     "that cannot be changed via the Windows API."),
    ("Timestamp Discrepancy Detection (Timestomping)",
     "If an attacker uses a tool to forge a file's creation date, they can only change "
     "the $SI timestamp. The $FN timestamp, written directly to the disk by the kernel, "
     "cannot be altered without directly writing to the raw disk. ShadowNet compares both "
     "timestamps. If the $SI date is backdated by more than 2 seconds vs. the $FN date, "
     "it fires a TIMESTOMP_CONFIRMED alert."),
    ("$USN Journal (Update Sequence Number)",
     "An indestructible log of every file operation on the volume. ShadowNet reads the USN "
     "journal and cross-references it with the MFT to detect ghost writes — cases where "
     "the disk changed but Windows has no record of authorizing the change. This is the "
     "fingerprint of a rootkit directly writing to raw hardware."),
    ("$LogFile",
     "The NTFS transaction log. Contains a record of in-progress disk operations. "
     "ShadowNet reads this to detect incomplete, forced, or corrupted transactions "
     "that often accompany aggressive anti-forensic tools."),
]
for name, desc in ntfs_points:
    p = doc.add_paragraph()
    p.add_run(f"{name}: ").bold = True
    p.add_run(desc)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)

# 4.3 Entropy Scanner
set_heading(doc, "4.3 Entropy Map Scanner (V5)", 2)
add_body(doc, "File: core/entropy_map_scanner.py  |  423 lines")
add_body(doc,
    "This module opens the same raw disk handle as the NTFS Parser and reads "
    "the volume cluster by cluster. For each cluster, it computes the Shannon Entropy of "
    "the binary data inside it. Entropy is a measure of how mathematically random the data is."
)
add_body(doc, "Why entropy matters for security:")
entropy_points = [
    "Normal computer files (Word documents, photos, text files) have entropy between 4.0 and 7.0. "
     "They contain repeating patterns and structure.",
    "Encrypted data (like that produced by ransomware) is mathematically indistinguishable from "
     "pure random noise. It scores an entropy value very close to 8.0 — the theoretical maximum for a byte.",
    "A data wiper that overwrites files with all-zeros (0x00) produces data with an entropy of 0.0 — "
     "the opposite extreme but equally unnatural.",
    "ShadowNet uses a Kolmogorov-Smirnov statistical test (K-S test) to check whether a cluster's "
     "byte distribution is statistically uniform (Entropy ~8.0) or statistically constant (Entropy ~0.0) "
     "and flags either as anomalous.",
    "If a contiguous run of more than 50,000 clusters (configurable) all show entropy above 7.85, "
     "this strongly suggests the presence of a hidden encrypted volume or active ransomware encryption."
]
for point in entropy_points:
    add_bullet(doc, point)

# 4.4 Behavioral Monitor  
set_heading(doc, "4.4 Real-time Behavioral Monitor (V5)", 2)
add_body(doc, "File: core/realtime_behavior_monitor.py  |  318 lines")
add_body(doc,
    "This module installs a low-level Windows keyboard hook (WH_KEYBOARD_LL) directly into "
    "the Windows message pump. This is the same level at which Windows kernel anti-cheat "
    "systems operate. Every key press on the machine is captured with nanosecond-precision "
    "timing (using Python's time.perf_counter_ns)."
)
add_body(doc, "Detection algorithm:")
algo_steps = [
    ("Inter-Keystroke Interval (IKI) Collection", 
     "The time gap between consecutive key presses is measured in milliseconds."),
    ("Statistical Feature Extraction",
     "The system computes: mean IKI, standard deviation (StDev), coefficient of variation (CV), "
     "and kurtosis of the timing distribution."),
    ("Gaussian Mixture Model (GMM) Analysis",
     "A GMM is fit to the IKI data to determine whether it is unimodal (consistent, bot-like) "
     "or bimodal (varied, human-like — humans naturally have a pause rhythm between words and characters)."),
    ("Rule-Based Verdict",
     "If CV < 0.15 (too consistent), StDev < 20ms (superhuman precision), or mean < 30ms "
     "(physically impossible for a human), the system fires a BOT_CONFIRMED alert."),
]
for name, desc in algo_steps:
    p = doc.add_paragraph()
    p.add_run(f"{name}: ").bold = True
    p.add_run(desc)
    p.paragraph_format.left_indent = Inches(0.3)

# 4.5 FIM
set_heading(doc, "4.5 File Integrity Monitor (V5)", 2)
add_body(doc, "File: core/file_integrity_monitor.py")
add_body(doc,
    "Uses the 'watchdog' Python library to install a real-time OS filesystem notification "
    "hook on configured directories. When any file matching a suspicious extension "
    "(.exe, .dll, .ps1, .bat) is created, modified, deleted, or moved, an alert is "
    "instantly fired. Deletion events are always flagged regardless of file type, "
    "as mass deletion is a common wiping technique."
)

# ─── 5. AI Analysis Core ───────────────────────────────────────────────────────
page_break(doc)  
set_heading(doc, "5. AI Analysis Core", 1, (0, 120, 212))

set_heading(doc, "5.1 Gemini Command Analyzer", 2)
add_body(doc, "File: core/gemini_command_analyzer.py")
add_body(doc,
    "When any module detects a suspicious command or process, the raw command-line string "
    "is sent to the Gemini 2.5 Flash language model with a detailed forensic analysis prompt. "
    "Gemini reads the intent of the command — not just the syntax — and returns a structured "
    "JSON verdict containing: threat category, severity level (CRITICAL/HIGH/MEDIUM/LOW), "
    "MITRE ATT&CK TTP codes, a plain-English explanation, and a recommended response action."
)
add_body(doc, "Before hitting the API, the analyzer applies:")
ai_features = [
    ("Command Decoder", "Detects and decodes Base64-encoded PowerShell commands before analysis."),
    ("Rate Limiter", "Enforces 20 API calls per minute to avoid quota exhaustion."),
    ("Circuit Breaker (3-state)", 
     "Tracks consecutive API failures. At 5 failures, it switches to OPEN state and routes "
     "all requests directly to the Keyword Fallback Engine for the next 5 minutes, "
     "then retries (HALF-OPEN state)."),
    ("Response Cache", "Caches identical command analysis results for 24 hours to save API calls."),
    ("Keyword Fallback Engine", 
     "When the AI is unavailable, this engine performs deterministic keyword matching. "
     "For V5 structural anomalies (NTFS, Entropy), it automatically assigns HIGH severity "
     "at 90% confidence, preserving all the raw evidence data in the incident report."),
]
for name, desc in ai_features:
    p = doc.add_paragraph()
    p.add_run(f"{name}: ").bold = True
    p.add_run(desc)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)

set_heading(doc, "5.2 Forensic Narrative Engine", 2)
add_body(doc, "File: core/forensic_narrative_engine.py")
add_body(doc,
    "Every 60 seconds, this module collects all findings from all four sensor modules "
    "simultaneously (NTFS, Entropy, Behavioral, Process). It synthesizes them into a "
    "single unified forensic narrative by sending the combined data as a structured prompt "
    "to Gemini. The output is a complete attack story — a timeline of what happened, which "
    "MITRE ATT&CK techniques were used, what evidence was destroyed, and what to investigate "
    "next. Reports are saved as both a JSON file and a human-readable Markdown file."
)
add_body(doc,
    "If Gemini is unavailable, a Rule-Based Fallback Synthesis generates the same structured "
    "report locally using deterministic logic, ensuring the narrative is never lost."
)

# ─── 6. Incident Response ─────────────────────────────────────────────────────
page_break(doc)
set_heading(doc, "6. Incident Response Pipeline", 1, (0, 120, 212))

set_heading(doc, "6.1 Emergency Snapshot Engine", 2)
add_body(doc, "File: core/emergency_snapshot.py")
add_body(doc,
    "The moment any detection fires, before the AI even begins its analysis, this engine "
    "captures a complete system snapshot in parallel using a 10-thread worker pool. "
    "Each snapshot collects: full Windows Event Logs, running process tree with all PIDs "
    "and command lines, active network connections, Volume Shadow Copy inventory, and "
    "filesystem metadata. All of this runs in under 100ms. "
    "If there is insufficient disk space or the capture fails, it retries up to 3 times "
    "with exponential backoff before giving up and logging the failure."
)

set_heading(doc, "6.2 Evidence Chain Signing", 2)
add_body(doc, "File: core/evidence_chain.py")
add_body(doc,
    "Every snapshot folder receives a cryptographic manifest. The SHA-256 hash of every "
    "file in the snapshot is computed and stored in a manifest.json file. At next startup, "
    "the system re-hashes every file and compares against the manifest. Any tampering "
    "with the captured evidence is immediately detected and reported."
)

set_heading(doc, "6.3 Structured Hash-Chain Logger", 2)
add_body(doc, "File: core/structured_logger.py")
add_body(doc,
    "Every forensic detection event is written to a tamper-evident log file "
    "(evidence/logs/forensic_findings.jsonl). Each log entry includes the SHA-256 hash "
    "of the previous entry, forming a blockchain-like chain. If any log entry is "
    "retroactively modified or deleted, the chain breaks at that exact point, "
    "proving exactly which record was tampered with. All logs are written in structured "
    "JSON Lines format for programmatic analysis."
)

set_heading(doc, "6.4 Incident Report Generator", 2)
add_body(doc, "File: core/incident_report_generator.py")
add_body(doc,
    "After the AI analysis completes, a complete Markdown incident report is generated "
    "and saved to evidence/incidents/<ID>/incident.md. The report includes: "
    "incident classification and severity, the exact triggering event (command or structural anomaly), "
    "process context (PID, user, parent process), AI threat assessment with MITRE codes, "
    "behavioral indicators, a list of all captured evidence, and a recommended response action."
)

# ─── 7. Resilience System ──────────────────────────────────────────────────────
page_break(doc)
set_heading(doc, "7. System Resilience", 1, (0, 120, 212))

set_heading(doc, "7.1 Thread Watchdog", 2)
add_body(doc, "File: core/watchdog.py")
add_body(doc,
    "All major detection modules (NTFS Parser, Entropy Scanner, Behavioral Monitor, "
    "Narrative Engine) run as separate background threads registered with the Watchdog. "
    "Each module sends a periodic heartbeat signal. The Watchdog checks every 5 seconds "
    "whether each thread is alive and whether its heartbeat has arrived within the timeout window."
)
add_body(doc, "Intelligent restart behavior:")
watchdog_points = [
    "If a thread dies or stops heartbeating, the Watchdog calls the registered factory function to create a clean replacement thread and start it.",
    "A crash-loop detection system tracks time since the last restart. If a thread dies again within 10 seconds of being restarted, it is likely crash-looping.",
    "A penalty of +2 is added to the restart counter, causing the exponential backoff timer to jump from seconds to minutes, protecting CPU from a spinning crash loop.",
    "Maximum backoff is capped at 120 seconds. After 10 total restarts, the module is marked as permanently dead and an alert is logged.",
    "The Dashboard UI reflects the real-time alive/dead status of each watchdog-registered module.",
]
for point in watchdog_points:
    add_bullet(doc, point)

set_heading(doc, "7.2 Alert Deduplication", 2)
add_body(doc, "File: core/alert_deduplication.py")
add_body(doc,
    "Every alert passes through a deduplication layer before creating an incident. "
    "A SHA-256 fingerprint is computed from the alert type and key data fields. "
    "If the same fingerprint was seen within the last 5 minutes, the alert is suppressed. "
    "This prevents a single event (like a process that executes 100 times rapidly) "
    "from flooding the system with 100 duplicate incidents. "
    "The deduplication cache is an LRU (Least Recently Used) ordered dictionary capped at 1000 entries."
)

# ─── 8. Visibility Layer ────────────────────────────────────────────────────────
page_break(doc)
set_heading(doc, "8. Visibility and Monitoring", 1, (0, 120, 212))

set_heading(doc, "8.1 Live Health Dashboard", 2)
add_body(doc, "File: core/dashboard.py  |  URL: http://127.0.0.1:8000")
add_body(doc,
    "A dark-mode web dashboard built with FastAPI and served locally. "
    "It auto-refreshes every 2 seconds and shows: system operational status, "
    "total detections count, system uptime, live CPU and RAM usage with visual bars, "
    "and — critically — the real-time active sub-systems list. "
    "This list is pulled directly from the Watchdog registry: "
    "only modules that are currently alive and registered appear as ACTIVE. "
    "If the NTFS Scanner is disabled in config or has crashed, it disappears from the list immediately."
)

set_heading(doc, "8.2 Health API", 2)
add_body(doc, "File: core/health_api.py  |  URL: http://127.0.0.1:8765/health")
add_body(doc,
    "A minimal JSON HTTP endpoint that external monitoring tools can poll. "
    "Returns the same structured health data as the dashboard in machine-readable format. "
    "Useful for integrating with uptime monitoring services or automated CI checks."
)

set_heading(doc, "8.3 SIEM Integration", 2)
add_body(doc, "File: core/siem_integration.py")
add_body(doc,
    "All alerts can be simultaneously forwarded to external Security Information and "
    "Event Management platforms. Supported targets: Splunk HTTP Event Collector (HEC), "
    "Elasticsearch, and any Syslog-compatible receiver. "
    "All three outputs are configurable via config_v5.yaml and are off by default."
)

# ─── 9. Configuration ────────────────────────────────────────────────────────
page_break(doc)
set_heading(doc, "9. Configuration System", 1, (0, 120, 212))

add_body(doc, "File: config/config_v5.yaml")
add_body(doc,
    "The entire system is controlled by a single YAML configuration file with two top-level "
    "sections: 'shadownet' (V4 core settings) and 'shadownet_v5' (V5 advanced modules)."
)

config_table = [
    ("enable_process_monitoring", "true/false", "Turns WMI + psutil process interception on or off"),
    ("enable_file_monitoring", "true/false", "Turns the filesystem watchdog monitor on or off"),
    ("suspicious_keywords", "list of strings", "Keywords that trigger an AI analysis request"),
    ("ntfs_parser.enabled", "true/false", "Requires Administrator. Enables raw $MFT scanning"),
    ("ntfs_parser.poll_interval_seconds", "integer (default 30)", "How often the MFT is swept for changes"),
    ("entropy_scanner.enabled", "true/false", "Requires Administrator. Enables cluster entropy scanning"),
    ("entropy_scanner.poll_interval_seconds", "integer (default 300)", "Scanning frequency in seconds"),
    ("entropy_scanner.hidden_volume_entropy_threshold", "float (default 7.85)", "Entropy level that triggers a hidden volume alert"),
    ("realtime_behavior.enabled", "true/false", "Turns the keyboard hook and GMM analysis on or off"),
    ("realtime_behavior.bot_confidence_alert_threshold", "float (default 0.65)", "Minimum confidence score to fire a bot alert"),
    ("forensic_narrative.synthesis_interval_seconds", "integer (default 60)", "How often the AI generates a unified attack narrative"),
    ("dashboard.enabled", "true/false", "Turns the web dashboard on or off"),
    ("dashboard.port", "integer (default 8000)", "Local port for the dashboard server"),
]
add_table(doc, ["Config Key", "Type / Default", "Description"], config_table)

# ─── 10. Technology Stack ───────────────────────────────────────────────────────
page_break(doc)
set_heading(doc, "10. Technology Stack", 1, (0, 120, 212))

stack_table = [
    ("Python 3.10+", "Primary language for all modules"),
    ("Google Gemini 2.5 Flash", "AI model for command intent analysis and narrative synthesis"),
    ("google-generativeai", "Official Python SDK for the Gemini API"),
    ("psutil", "Cross-platform process and system resource inspection"),
    ("numpy / scipy", "Entropy computation, K-S statistical tests, IKI analysis"),
    ("scikit-learn (sklearn)", "Gaussian Mixture Model fitting for behavioral analysis"),
    ("pywin32 / wmi", "WMI event subscriptions for real-time process creation notifications"),
    ("ctypes / ctypes.windll", "Raw Windows API calls: CreateFileW (disk handle), keyboard hooks, WM_QUIT dispatch"),
    ("watchdog (pip package)", "Cross-platform filesystem event notifications for FIM"),
    ("FastAPI + uvicorn", "Live health dashboard web server"),
    ("PyYAML", "Configuration file parsing"),
    ("python-dotenv", "Secure API key loading from .env file"),
    ("threading / queue", "Multi-threaded concurrent module execution"),
    ("hashlib (SHA-256)", "Evidence integrity hashing and log chain computation"),
]
add_table(doc, ["Technology", "Purpose"], stack_table)

# ─── 11. File Structure ─────────────────────────────────────────────────────────
page_break(doc)
set_heading(doc, "11. Project File Structure", 1, (0, 120, 212))

structure = [
    ("shadownet_nexus_v5.py", "Main entry point. Orchestrates all modules and the main loop."),
    ("config/config_v5.yaml", "Unified configuration file for all V4 and V5 settings."),
    ("core/ntfs_artifact_parser.py", "Raw NTFS $MFT, $USN Journal, $LogFile parser. 774 lines."),
    ("core/entropy_map_scanner.py", "Cluster-level entropy map scanner for hidden/wiped data. 423 lines."),
    ("core/realtime_behavior_monitor.py", "Windows keyboard hook + GMM bot detection. 318 lines."),
    ("core/process_monitor.py", "WMI + psutil dual-mode process interception. 251 lines."),
    ("core/file_integrity_monitor.py", "Cross-platform filesystem change monitoring."),
    ("core/gemini_command_analyzer.py", "Gemini AI command intent analyzer with circuit breaker + fallback."),
    ("core/forensic_narrative_engine.py", "Periodic synthesis of all findings into a unified narrative."),
    ("core/incident_report_generator.py", "Generates structured Markdown incident reports."),
    ("core/emergency_snapshot.py", "Parallel system state capture triggered on every detection."),
    ("core/evidence_chain.py", "Cryptographic SHA-256 manifest signing and verification."),
    ("core/structured_logger.py", "SHA-256 hash-chain forensic log handler."),
    ("core/watchdog.py", "Thread registry with intelligent crash-loop detection and auto-restart."),
    ("core/alert_deduplication.py", "LRU fingerprint-based alert deduplication cache."),
    ("core/dashboard.py", "FastAPI live health dashboard (http://127.0.0.1:8000)."),
    ("core/alert_manager.py", "Unified alert routing to Console, Slack, Email, and SIEM."),
    ("core/behavior_monitor.py", "V4 behavioral simulation engine."),
    ("core/proactive_evidence_collector.py", "Orchestrates snapshot capture and evidence vault management."),
    ("tests/live_combat_test.py", "Simulates real attacks to verify all detection modules work live."),
    ("evidence/", "Output directory for all snapshots, logs, reports, and artifacts."),
    ("prompts/enhanced_prompts.py", "All Gemini AI prompt templates."),
    ("utils/", "Helper utilities: model selector, OS detector, cache manager, command decoder."),
]
add_table(doc, ["File / Directory", "Purpose"], structure)

# ─── 12. Key Claims ────────────────────────────────────────────────────────────
page_break(doc)
set_heading(doc, "12. Key Technical Claims", 1, (0, 120, 212))

add_body(doc,
    "The following capabilities represent genuinely novel combinations of techniques "
    "not simultaneously available in any known open-source or commercial tool as of Q1 2026:"
)
claims = [
    "Live $MFT binary parsing without requiring a disk image, running against a mounted, active Windows volume.",
    "Simultaneous cluster-level entropy mapping with K-S uniformity testing on a live system.",
    "Detection of wipe pattern signatures (near-zero entropy) in unallocated disk slack space.",
    "Real OS-level keyboard hook for keystroke dynamics via WH_KEYBOARD_LL with nanosecond precision timestamps.",
    "GMM-based bimodal distribution analysis to distinguish human from automated input.",
    "Cross-module LLM forensic narrative synthesis that correlates NTFS, entropy, behavioral, and process findings into a single attack story.",
    "Circuit breaker + rate limiter + keyword fallback architecture ensuring incident reports are always generated regardless of API availability.",
    "Hash-chain tamper-evident logging that proves retroactive log modification at entry-level granularity.",
    "Watchdog-driven thread resurrection with crash-loop penalty backoff scaling from seconds to minutes.",
    "Real-time dashboard that reflects true module status from the Watchdog registry, not from a hardcoded list.",
]
for i, claim in enumerate(claims, 1):
    p = doc.add_paragraph()
    p.add_run(f"{i}. ").bold = True
    p.add_run(claim)
    p.paragraph_format.space_after = Pt(3)

# ─── Save ──────────────────────────────────────────────────────────────────────
output_path = "ShadowNet_Nexus_v5_Documentation.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
